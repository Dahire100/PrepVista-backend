import gc
import time
import google.generativeai as genai
import pdfplumber
import threading
import json
import re
import os
import tempfile
import sys
import subprocess
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from collections import defaultdict
import random
from datetime import datetime
# Imports for Word Document Generation
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
# Core Functionality
import speech_recognition as sr
from gtts import gTTS
import pygame
import cv2
# REMOVED tkinter and all GUI-related imports and code for backend compatibility
# from PIL import Image, ImageTk

# --- 2. CONFIGURATION ---
API_KEY = os.getenv('GEMINI_API_KEY_1', 'your_gemini_api_key_here')# Configure the API
model = None
try:
    if not API_KEY or "YOUR_VALID_API_KEY_HERE" in API_KEY:
        print("🛑 ERROR: Please replace 'YOUR_VALID_API_KEY_HERE' with your actual Gemini API key.")
    else:
        genai.configure(api_key=API_KEY)
        model = genai.GenerativeModel('gemini-2.5-flash')
        print("✅ Gemini API Key configured and model initialized.")
except Exception as e:
    print(f"🛑 ERROR: Could not configure Gemini API. Please check your key. Error: {e}")
    print("⚠  Continuing without API access - some features will not work")

# --- 3. DATA STRUCTURES ---
@dataclass
class Question:
    type: str
    question: str
    category: str = "general"  # More specific categorization
@dataclass
class Evaluation:
    score: int
    feedback: str
    suggestions: str
    corrected_answer: str
    keywords_matched: List[str] = None
@dataclass
class InterviewPlan:
    branch: str
    skills_summary: str
    projects_summary: str
    question_bank: List[Question]
    resume_keywords: List[str] = None
@dataclass
class InterviewResult:
    question: Question
    user_answer: str
    evaluation: Evaluation
    face_detected: bool = True
        candidate_name: str = "Candidate"  # Optional parameter for backward compatibility

# Initialize pygame for audio playback with better error handling
try:
    pygame.mixer.init(frequency=22050, size=-16, channels=2, buffer=512)
    print("✅ Pygame mixer initialized successfully.")
except pygame.error as e:
    print(f"🛑 Pygame mixer could not be initialized. Audio playback will be disabled. Error: {e}")

# Initialize speech recognition with better configuration
recognizer = sr.Recognizer()
recognizer.energy_threshold = 300
recognizer.pause_threshold = 1.5
recognizer.dynamic_energy_threshold = True
microphone = None
try:
    microphone = sr.Microphone()
    print("Calibrating microphone for ambient noise...")
    with microphone as source:
        recognizer.adjust_for_ambient_noise(source, duration=3)
        
    print("✅ Microphone calibration complete.")
except Exception as e:
    print(f"🛑 Microphone not found or could not be initialized. Voice input will not work. Error: {e}")

# --- 4. HELPER FUNCTIONS ---

def get_info_from_gemini(prompt: str, is_json_response: bool = True, max_retries: int = 3) -> Optional[Any]:
    """Get and parse responses from the Gemini API with retry logic."""
    if not model:
        print("Gemini model not available, using fallback.")
        if "interview plan" in prompt:
            return get_fallback_interview_plan()
        return None
    
    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt)
            text = response.text.strip()
            
            if not is_json_response:
                return text
            
            # Try to extract JSON from response
            json_match = re.search(r'```json\s*(.+?)\s*```', text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1)
            elif text.startswith('{') and text.endswith('}'):
                json_text = text
            else:
                # Attempt to find the JSON block if formatting is off
                start_index = text.find('{')
                end_index = text.rfind('}') + 1
                if start_index != -1 and end_index != 0:
                    json_text = text[start_index:end_index]
                else:
                    raise ValueError("No valid JSON object found in the response.")
            
            return json.loads(json_text)
        except Exception as e:
            print(f"AI Response Error (Attempt {attempt+1}/{max_retries}): {e}")
            if attempt == max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
            else:
                return None

def get_fallback_interview_plan():
    """Fallback interview plan when Gemini API is not available."""
    return {
        "branch": "Computer Science (Fallback)",
        "skills_summary": "Programming, Problem Solving",
        "projects_summary": "Software development projects",
        "question_bank": [
            {"type": "resume", "question": "Tell me about your most significant project."},
            {"type": "core", "question": "Explain a core concept from Data Structures."},
            {"type": "core", "question": "Write a simple SQL query to find all users from a 'users' table in New York."},
            {"type": "HR", "question": "Where do you see yourself in 5 years?"}
        ]
    }

def parse_resume(pdf_path: str) -> Optional[str]:
    """Extracts text from a PDF using pdfplumber."""
    print("Parsing resume...")
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
        
        if not text.strip():
            print("Warning: Could not extract meaningful text from the PDF.")
            return None
        
        print("Resume parsed successfully.")
        return text
    except Exception as e:
        print(f"ERROR: Failed to parse PDF file. {e}")
        return None

def extract_keywords_from_resume(resume_text: str, max_keywords: int = 15) -> List[str]:
    """Extract important keywords from resume using AI."""
    prompt = f"""Analyze the following resume text and extract the {max_keywords} most important technical skills, technologies, and domain-specific keywords.
Return ONLY a JSON array of strings.

Resume Text:
---
{resume_text[:3000]}
---"""
    
    keywords = get_info_from_gemini(prompt)
    if keywords and isinstance(keywords, list):
        return keywords[:max_keywords]
    
    # Fallback to simple keyword extraction
    technical_terms = [
        "python", "java", "javascript", "sql", "html", "css", "react", "angular", "vue",
        "node", "express", "django", "flask", "aws", "azure", "gcp", "docker", "kubernetes",
        "machine learning", "ai", "data science", "analytics", "backend", "frontend",
        "fullstack", "devops", "ci/cd", "agile", "scrum", "rest", "api", "graphql", "microservices"
    ]
    found_keywords = []
    for term in technical_terms:
        if term in resume_text.lower():
            found_keywords.append(term)
    return found_keywords[:max_keywords]

def get_branch_from_resume(resume_text: str) -> str:
    """Uses AI to determine the candidate's academic/engineering branch."""
    prompt = f"""Analyze the following resume text and identify the candidate's primary academic or engineering branch.
Respond with ONLY the name of the branch (e.g., 'Computer Science', 'Mechanical Engineering', 'Business Administration').

Resume Text:
---
{resume_text[:2000]}
---"""
    
    branch = get_info_from_gemini(prompt, is_json_response=False)
    return branch or "General"

def generate_interview_plan(resume_text: str, previous_questions: List[str] = None) -> Optional[InterviewPlan]:
    """Uses AI to create a complete, structured interview plan tailored to the candidate's branch."""
    print("Determining candidate's branch from resume...")
    branch_name = get_branch_from_resume(resume_text) or "Unknown"
    print(f"Detected Branch: {branch_name}")
    
    resume_keywords = extract_keywords_from_resume(resume_text) or []
    print(f"Extracted Keywords: {', '.join(resume_keywords) if resume_keywords else 'None'}")
    
    # Build prompt for AI
    prompt = f"""You are an expert technical interviewer. Analyze the following resume and create a comprehensive interview plan.

Resume Text:
---
{resume_text}
---

Candidate's Field/Branch: {branch_name}
Key Skills/Technologies: {', '.join(resume_keywords)}

Generate an interview plan with exactly 10 diverse questions. Return a JSON object with this structure:
{{
  "branch": "{branch_name}",
  "skills_summary": "<concise summary of technical skills>",
  "projects_summary": "<summary of projects and experience>",
  "question_bank": [
    {{
      "type": "technical|behavioral|project|problem-solving",
      "question": "<question text>",
      "category": "<specific category>"
    }}
  ]
}}

Question distribution:
- 4 technical questions based on skills mentioned
- 3 project-based questions
- 2 behavioral questions
- 1 problem-solving/coding question

Return ONLY valid JSON."""
    
    print("Generating interview plan with AI...")
    try:
        ai_response = get_info_from_gemini(prompt)
        
        if not ai_response:
            print("Using fallback interview plan")
            ai_response = get_fallback_interview_plan()
        
        # Convert questions to Question objects
        questions = [
            Question(
                type=q.get('type', 'general'),
                question=q.get('question', ''),
                category=q.get('category', 'general')
            )
            for q in ai_response.get('question_bank', [])
        ]
        
        plan = InterviewPlan(
            branch=ai_response.get('branch', branch_name),
            skills_summary=ai_response.get('skills_summary', ''),
            projects_summary=ai_response.get('projects_summary', ''),
            question_bank=questions,
            resume_keywords=resume_keywords
        )
        
        print(f"Interview plan generated successfully with {len(questions)} questions.")
        return plan
    
    except Exception as e:
        print(f"Error generating interview plan: {e}")
        return None

def evaluate_answer(question: Question, answer: str, resume_keywords: List[str] = None) -> Evaluation:
    """Evaluate interview answer using AI."""
    try:
        prompt = f"""You are an expert technical interviewer. Evaluate this interview answer.

Question Type: {question.type}
Question: {question.question}
Candidate's Answer: {answer}

Provide evaluation in JSON format:
{{
  "score": <0-10>,
  "feedback": "<positive feedback>",
  "suggestions": "<specific improvements>",
  "corrected_answer": "<ideal answer example>",
  "keywords_matched": [<relevant keywords>]
}}

Return ONLY valid JSON."""
        
        result = get_info_from_gemini(prompt)
        
        if result:
            return Evaluation(
                score=result.get('score', 5),
                feedback=result.get('feedback', 'Answer received'),
                suggestions=result.get('suggestions', 'Keep practicing'),
                corrected_answer=result.get('corrected_answer', ''),
                keywords_matched=result.get('keywords_matched', [])
            )
    except Exception as e:
        print(f"Error evaluating answer: {e}")
    
    # Fallback evaluation
    return Evaluation(
        score=7,
        feedback="Good answer",
        suggestions="Consider adding more specific examples",
        corrected_answer="",
        keywords_matched=[]
    )

def add_formatted_paragraph(document, text: str, bold: bool = False, italic: bool = False, color: Tuple[int, int, int] = None):
    """Add a formatted paragraph to the document."""
    para = document.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def generate_word_report(interview_plan: InterviewPlan, interview_results: List[InterviewResult], camera_verified: bool = True) -> Tuple[str, float]:
    """Generate comprehensive Word report."""
    try:
        document = docx.Document()
        
        # Title
        title = document.add_heading('AI Interview Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calculate average score
        if interview_results:
            avg_score = sum(r.evaluation.score for r in interview_results) / len(interview_results)
        else:
            avg_score = 0
        
        # Overview Section
        document.add_heading('Overview', level=1)
        table = document.add_table(rows=0, cols=2)
        table.style = 'Light Shading-Accent 1'
        
        overview_data = [
            ("Candidate Branch", interview_plan.branch),
            ("Skills Focus", interview_plan.skills_summary),
            ("Projects Focus", interview_plan.projects_summary),
            ("Average Score", f"{avg_score:.1f}/10"),
            ("Camera Verification", "Passed" if camera_verified else "Failed"),
            ("Total Questions", len(interview_results))
        ]
        
        for key, value in overview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            row_cells[0].paragraphs[0].runs[0].bold = True
        
        document.add_paragraph()
        
        # Detailed Questions
        document.add_heading('Detailed Question Analysis', level=1)
        
        for i, result in enumerate(interview_results):
            document.add_heading(f"Question {i+1}: {result.question.type.capitalize()} Question", level=2)
            
            add_formatted_paragraph(document, "Question:", bold=True)
            add_formatted_paragraph(document, result.question.question, italic=True)
            
            add_formatted_paragraph(document, "Candidate's Answer:", bold=True)
            add_formatted_paragraph(document, result.user_answer)
            
            score_color = (0, 128, 0) if result.evaluation.score >= 8 else (0, 102, 204) if result.evaluation.score >= 6 else (204, 0, 0)
            add_formatted_paragraph(document, "Score:", bold=True)
            add_formatted_paragraph(document, f"{result.evaluation.score}/10", color=score_color, bold=True)
            
            add_formatted_paragraph(document, "Feedback:", bold=True)
            add_formatted_paragraph(document, result.evaluation.feedback)
            
            add_formatted_paragraph(document, "Suggestions for Improvement:", bold=True)
            add_formatted_paragraph(document, result.evaluation.suggestions)
            
            document.add_paragraph()  # spacing
        
        # Save document
        report_filename = f"AI_Interview_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(report_filename)
        print(f"Word report successfully generated: {report_filename}")
        return report_filename, avg_score
    
    except Exception as e:
        print(f"Error generating report: {e}")
        return None, 0
# --- 7. MAIN EXECUTION ---
def main():
    """Main entry point for the backend application (for Render deployment)."""
    print("Backend logic loaded for API deployment. GUI code removed.")
    # Instead of starting Tkinter, this should wire up to Flask/FastAPI endpoints
    # Example:
    # from flask import Flask
    # app = Flask(__name__)
    # ...
    # app.run()

if __name__ == "__main__":
    main()
